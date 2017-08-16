# coding:utf-8

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextContainer
from pdfminer.pdfpage import PDFPage
from operator import itemgetter

import docx

import subprocess
import os
import heapq

import re
import redis

import jieba.posseg as pseg  # 需要另外加载一个词性标注模块
import spacy

from redismap import RedisMap

import logging
LOG_FORMAT = '%(asctime)s - %(filename)s[line:%(lineno)d] - %(threadName)s - %(levelname)s: %(message)s'
logging.basicConfig(level=logging.INFO, format=LOG_FORMAT)


r = redis.Redis(host='127.0.0.1', port=6379)


def normal_weight(weight_paras_map, nmerge=1):
    """
    归一化weight为整数，同时可选合并同一weight下的段落。
    @param weight_paras_map: 源字典，需要为RedisMap对象，key为weight， value为段落列表
    @param nmerge: 将同一weight下的段落按照参数nmerge个合并为一个以便增加每一次结巴处理的字串长度。减少循环次数。
    """
    base_weight = min(map(float, weight_paras_map.keys()))  # 找出基准(最小)权重, redis取出的数全是str,先全部强转成float

    for w in weight_paras_map.keys():
        weight = int((float(w) / base_weight - 1) * 10) + 1
        logging.debug('rename key : %s -> %s' % (w, str(weight)))

        if nmerge > 1:
            if weight not in weight_paras_map:
                weight_paras_map[weight] = []

            ps = weight_paras_map[w]
            new_ps = weight_paras_map[weight]

            i = 0
            length = len(ps)
            while i < length:
                new_ps.append('\n'.join(ps[i:i+nmerge]))
                i += nmerge

            del weight_paras_map[w]
        else:
            weight_paras_map.rename_key(w, weight)


def convert_pdf_2_text(path):

    weight_category = RedisMap(r, 'pdf_paras')
    weight_category.clear()     # 先清除之前的记录

    rsrcmgr = PDFResourceManager()
    device = PDFPageAggregator(rsrcmgr, laparams=LAParams())
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    with open(path, 'rb') as fp:

        def get_line_numbers(s):
            return len(re.findall(r'\n', s, re.DOTALL))

        for page in PDFPage.get_pages(fp, set()):
            interpreter.process_page(page)
            layouts = device.get_result()
            for x in layouts:
                # logging.debug(type(x))
                if isinstance(x, LTTextContainer):
                    text = x.get_text()
                    # logging.debug('BEGIN %s:' % type(x), x.height / get_line_numbers(text), text)
                    line_height = int((x.height / get_line_numbers(text)) * 1000) / 1000.0  # 保留三位小数

                    if line_height in weight_category:
                        weight_category[line_height].append(text)
                    else:
                        weight_category[line_height] = [text]

    device.close()

    normal_weight(weight_category, 10)

    return weight_category


def convert_doc_2_docx(docName):

    if docName.endswith('.doc'):
        docx_name = docName + 'x'
        if os.path.exists(docx_name):
            return True

        if os.path.exists(docName):
            dir_name = os.path.dirname(docName)
            if dir_name == '':
                dir_name = '.'

            subprocess.call(['soffice', '--headless', '--convert-to', 'docx', '--outdir', dir_name, docName])
            if os.path.exists(docx_name):   # xxx.doc -> xxx.docx
                os.remove(docName)
                return True

    return False


def convert_docx_2_text(docName):
    weight_category = RedisMap(r, 'doc_paras')
    weight_category.clear()     # 使用前先清空之前的记录

    doc = docx.Document(docName)
    paras = doc.paragraphs  # 添加文档中的段落

    def extend_paras_from_table(table):
        for i in range(len(table.rows)):
            for j in range(table._column_count):
                cell = table.cell(i, j)

                tc = cell._tc
                if tc.left != j or tc.top != i:  # 忽略融合的单元格
                    continue

                paras.extend(cell.paragraphs)
                for t in cell.tables:
                    extend_paras_from_table(t)  # 添加表格单元中的嵌套表格

    for table in doc.tables:  # 添加表格中的段落
        extend_paras_from_table(table)

    def get_size(style):  # 获得字号
        if style.rPr and style.rPr.sz_val:
            return style.rPr.sz_val / 6350
        return get_size(style.base_style)

    def is_bold(style):  # 是否是粗体
        if style.rPr and style.rPr.b:
            return style.rPr.b.val
        return is_bold(style.base_style) if style.base_style else False

    for p in paras:
        style = doc.styles.element.get_by_id(p.style.style_id)  # 此处style类型与'p.style'的类型不一致,此处Style包含更多信息。

        base_size = get_size(style)
        base_bold = is_bold(style)

        for run in p.runs:  # 处理同一段中不同样式的文字
            if len(run.text.strip()) < 1:
                continue

            size = base_size
            bold = base_bold

            if run._r and run._r.rPr:  # 如果有内嵌的字体大小/粗体等控制，则更新相应的值。
                if run._r.rPr.sz_val:
                    size = run._r.rPr.sz_val / 6350
                if run._r.rPr.b is not None:  # r._r.rPr.b的‘b’所属类定制了bool类型，因此必须使用‘is not None’判断！
                    bold = run._r.rPr.b.val

            if bold:
                size *= 1.5

            text = run.text
            if size in weight_category:
                weight_category[size].append(text)
            else:
                weight_category[size] = [text]

    normal_weight(weight_category, 10)

    return weight_category


def english_process(nlp, article, weight, word_freq):
    doc = nlp(article)
    # for word in doc.noun_chunks:  # 名词短语
    # for word in doc.ents:         # 实体
    for token in doc:
        if not token.tag_.startswith('NN'):  # 只取名词
            continue

        word = token.lemma_  # 词干化
        if word in word_freq:
            word_freq[word] += 1 * weight
        else:
            word_freq[word] = 1 * weight


def jieba_process(weight_paras):
    word_freq = {}
    words_count = 0
    eng_count = 0
    nlp = spacy.load('en_default')  # 使用spacy处理英文

    for weight, ps in weight_paras.items():
        weight = int(weight)
        logging.info('process weight-%d, length=%d' % (weight, len(ps)))

        for p in ps:
            logging.debug('process :%s' % p)
            words = pseg.cut(p)     # 使用Jieba的词性标注模块进行处理
            engs = []               # 统计出段落中的英文，进行后续的spacy处理

            for word in words:
                flag = word.flag
                word = word.word
                if not flag.startswith('n') and not flag == 'eng':  # 只处理名词词性与英语词汇
                    continue

                words_count += 1

                if flag == 'eng':   # 若是英语，则添加到eng列表，待后续处理
                    eng_count += 1
                    engs.append(word)
                else:               # 中文根据权重进行累加词频
                    if word in word_freq:
                        word_freq[word] += 1 * weight
                    else:
                        word_freq[word] = 1 * weight

            english_process(nlp, u' '.join(engs), weight, word_freq)

    freq_word = []
    for word, freq in word_freq.items():
        freq_word.append((word, freq))

    max_number = 20

    for word, freq in heapq.nlargest(max_number, freq_word, key=itemgetter(1)):
        logging.info('%s %s' % (word, freq))


def process(filename):
    weight_paras = None
    if filename.endswith('.pdf'):
        weight_paras = convert_pdf_2_text(filename)
    elif filename.endswith('.docx'):
        weight_paras = convert_docx_2_text(filename)
    elif filename.endswith('.doc'):
        if convert_doc_2_docx(filename):
            weight_paras = convert_docx_2_text(filename + 'x')

    if weight_paras:
        jieba_process(weight_paras)
        weight_paras.clear()

if __name__ == '__main__':
    process('data/TN_POLADROID_DEMO_APP_v0.1.docx')

    logging.info('分割线----------------------------')

    process('data/TN_INTUITION_ENGINE_v0.1.doc')

    # process('data/TN_POLADROID_DEMO_APP_v0.1.pdf')

